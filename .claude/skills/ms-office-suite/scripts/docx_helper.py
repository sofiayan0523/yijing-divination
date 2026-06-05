#!/usr/bin/env python3
"""Word document helper — create, read, and edit .docx files.

Usage:
    python docx_helper.py create <output.docx> <title>
    python docx_helper.py read <input.docx>
    python docx_helper.py info <input.docx>
"""
import sys
import json


def create_document(output_path: str, title: str) -> None:
    """Create a simple document with a title as a starting point."""
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    doc.add_heading(title, level=0)
    doc.add_paragraph("")  # Empty paragraph as spacer

    doc.save(output_path)
    print(f"Created: {output_path}")


def read_document(input_path: str) -> None:
    """Extract text content from a document."""
    from docx import Document

    doc = Document(input_path)
    for para in doc.paragraphs:
        if para.text.strip():
            style = para.style.name if para.style else "Normal"
            print(f"[{style}] {para.text}")

    for i, table in enumerate(doc.tables):
        print(f"\n--- Table {i + 1} ---")
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            print(" | ".join(cells))


def get_info(input_path: str) -> None:
    """Get document metadata as JSON."""
    from docx import Document

    doc = Document(input_path)
    props = doc.core_properties
    info = {
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "title": props.title or "",
        "author": props.author or "",
        "styles": [s.name for s in doc.styles if s.type is not None][:20],
    }
    print(json.dumps(info, indent=2))


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)

    cmd = sys.argv[1]
    if cmd == "create":
        create_document(
            sys.argv[2],
            sys.argv[3] if len(sys.argv) > 3 else "Untitled Document",
        )
    elif cmd == "read":
        read_document(sys.argv[2])
    elif cmd == "info":
        get_info(sys.argv[2])
    else:
        print(f"Unknown command: {cmd}")
        sys.exit(1)
